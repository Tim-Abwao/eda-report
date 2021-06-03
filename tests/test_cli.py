import subprocess
import unittest
from pathlib import Path

import pandas as pd
from docx import Document


class TestCLI(unittest.TestCase):
    def setUp(self):
        # Create a dummy input file
        pd.DataFrame(range(50)).to_csv("test-data.csv", index=False)

    def test_cli_with_all_args(self):
        # Run the program through the command-line interface
        subprocess.run(
            [
                "python",
                "-m",
                "eda_report",
                "test-data.csv",
                "-o",
                "cli-test-1.docx",
                "-t",
                "CLI Test",
                "-T",
                "0",
            ]
        )
        # Check if a report is generated as specified
        self.assertTrue(Path("cli-test-1.docx").is_file())
        self.assertTrue(
            # Check if the title is as specified
            Document("cli-test-1.docx").paragraphs[0].text,
            "CLI Test",
        )

    def test_cli_with_default_args(self):
        # Run the program through the command-line interface
        subprocess.run(
            [
                "python",
                "-m",
                "eda_report",
                "test-data.csv",
            ]
        )
        # Check if a report is generated with the default output file-name
        self.assertTrue(Path("eda-report.docx").is_file())
        self.assertTrue(
            # Check if the default title was set
            Document("eda-report.docx").paragraphs[0].text,
            "Exploratory Data Analysis Report",
        )

    def tearDown(self):
        for filename in {
            "test-data.csv",
            "cli-test-1.docx",
            "eda-report.docx",
        }:
            file = Path(filename)
            if file.is_file():
                file.unlink()  # Delete the file
