import unittest
from pathlib import Path

import pandas as pd
import seaborn as sns
from docx import Document
from eda_report import get_word_report


class TestDocumentProperties(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.outfile = "test-report.docx"

    @classmethod
    def tearDownClass(cls):
        Path(cls.outfile).unlink()  # Delete the file
        del cls.outfile

    def test_report_creation_in_multivariate_mixed_type_data(self):
        # Generate the report
        get_word_report(
            data=sns.load_dataset("iris"),
            output_filename=self.outfile,
            title="Mutlivariate",
        )
        # Check if the report is actually created and saved
        self.assertTrue(Path(self.outfile).is_file())
        # Check if the title in the document is as stipulated
        self.assertEqual(
            Document(self.outfile).paragraphs[0].text, "Mutlivariate"
        )

    def test_report_creation_in_numeric_univariate_data(self):
        # Generate the report
        get_word_report(
            data=pd.DataFrame(range(50)),
            output_filename=self.outfile,
            title="Numeric Univariate",
        )
        # Check if the report is actually created and saved
        self.assertTrue(Path(self.outfile).is_file())
        # Check if the title in the document is as stipulated
        self.assertEqual(
            Document(self.outfile).paragraphs[0].text,
            "Numeric Univariate",
        )

    def test_report_creation_in_categorical_univariate_data(self):
        # Generate the report
        get_word_report(
            data=pd.DataFrame(["only one item"]),
            output_filename=self.outfile,
            title="Categorical Univariate",
        )
        # Check if the report is actually created and saved
        self.assertTrue(Path(self.outfile).is_file())
        # Check if the title in the document is as stipulated
        self.assertEqual(
            Document(self.outfile).paragraphs[0].text,
            "Categorical Univariate",
        )
