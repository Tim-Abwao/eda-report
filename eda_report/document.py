import logging
from typing import Iterable, Sequence, Union

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from pandas import DataFrame, Series

from eda_report._content import _ReportContent

logging.basicConfig(
    format="[%(levelname)s %(asctime)s.%(msecs)03d] %(message)s",
    level=logging.INFO,
    datefmt="%H:%M:%S",
)
# Set matplotlib logging level to WARNING.
mpl_logger = logging.getLogger("matplotlib")
mpl_logger.setLevel(logging.WARNING)


class ReportDocument(_ReportContent):
    """Creates a report :class:`~docx.document.Document` with analysis results.

    The report consists of 3 main sections:

    #. An **Overview** of the data and its features.
    #. **Univariate Analysis**: Summary statistics and graphs for each feature.
    #. **Bivariate Analysis**: Pair-wise comparisons of numerical features.

    Args:
        data (Iterable): The data to analyze.
        title (str, optional): The title to assign the report. Defaults to
            "Exploratory Data Analysis Report".
        graph_color (str, optional): The color to apply to the graphs.
            Defaults to "cyan".
        groupby_variable (Union[str, int], optional): The column to
            use to group values. Defaults to None.
        output_filename (str, optional): The name/path to save the document
            to. Defaults to "eda-report.docx".
        table_style (str, optional): The style to apply to the tables created.
            Defaults to "Table Grid".
    """

    def __init__(
        self,
        data: Iterable,
        *,
        title: str = "Exploratory Data Analysis Report",
        graph_color: str = "cyan",
        groupby_variable: Union[str, int] = None,
        output_filename: str = "eda-report.docx",
        table_style: str = "Table Grid",
    ) -> None:
        super().__init__(
            data,
            title=title,
            graph_color=graph_color,
            groupby_variable=groupby_variable,
        )
        self.OUTPUT_FILENAME = output_filename
        self.TABLE_STYLE = table_style
        self.document = Document()  # Initialize report document
        self._create_cover_page()
        self._get_univariate_analysis()

        if self.dataset._correlation_values is not None:
            self._get_bivariate_analysis()

        self._to_file()
        logging.info(f"Done. Results saved as {self.OUTPUT_FILENAME!r}")

    def _create_cover_page(self) -> None:
        """Add a title and overview of the data."""
        self.document.add_heading(self.TITLE, level=0)
        self.document.add_paragraph(self.intro_text)
        self._get_numeric_overview_table()
        self._get_categorical_overview_table()
        self.document.add_page_break()

    def _get_numeric_overview_table(self) -> None:
        """Create a table with an overview of the numeric features present."""
        if self.dataset._numeric_stats is None:
            return None
        else:
            heading = self.document.add_heading(
                "Overview of Numeric Features", level=1
            )
            self._format_paragraph_spacing(heading)
            # count | avg | stddev | min | 25% | 50% | 75% | max
            self._create_table(
                data=self.dataset._numeric_stats,
                header=True,
                column_widths=(1.2,) + (0.7,) * 8,
                font_size=8.5,
                style="Normal Table",
            )

    def _get_categorical_overview_table(self) -> None:
        """Create a table with an overview of the categorical features
        present.
        """
        if self.dataset._categorical_stats is None:
            return None
        else:
            heading = self.document.add_heading(
                "Overview of Categorical Features", level=1
            )
            self._format_paragraph_spacing(heading)
            # column-name | count | unique | top | freq | relative freq
            self._create_table(
                data=self.dataset._categorical_stats,
                header=True,
                column_widths=(1.2,) + (0.9,) * 5,
                font_size=8.5,
                style="Normal Table",
            )

    def _get_univariate_analysis(self) -> None:
        """Get a brief introduction, summary statistics, and graphs for each
        individual variable.
        """
        univariate_heading = self.document.add_heading(
            "1. Univariate Analysis", level=1
        )
        self._format_paragraph_spacing(univariate_heading, before=0, after=0)
        for idx, variable in enumerate(self.variables.values(), start=1):
            var_name = variable.name
            description = self.variable_descriptions[var_name]
            summary_stats = Series(self.univariate_stats[var_name]).to_frame()
            graphs = self.univariate_graphs[var_name]
            contingency_table = self.contingency_tables.get(var_name)
            normality_tests = self.normality_tests.get(var_name)
            # Variable's title and brief description
            heading = self.document.add_heading(
                f"1.{idx} {var_name}".title(), level=2
            )
            self._format_paragraph_spacing(heading, before=12, after=5)
            self.document.add_paragraph(description)
            # Summary statistics table
            stats_heading = self.document.add_heading(
                "Summary Statistics", level=4
            )
            stats_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._create_table(summary_stats, column_widths=[2.5, 2])
            # Images of plotted graphs
            for name, image in graphs.items():
                width = 3.3 if name == "prob_plot" else 4.2
                self.document.add_picture(image, width=Inches(width))
                picture_paragraph = self.document.paragraphs[-1]
                picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if contingency_table is not None:
                contingency_table_heading = self.document.add_heading(
                    "Contingency table", level=4
                )
                contingency_table_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                context = self.document.add_paragraph(
                    f"Index = '{var_name}', "
                    f"Columns = '{self.GROUPBY_DATA.name}' "
                )
                context.alignment = WD_ALIGN_PARAGRAPH.CENTER
                context.runs[0].font.size = Pt(8)
                n_cols = contingency_table.shape[1]
                max_width = 5.2 if n_cols > 5 else 3.2
                col_width = max_width / n_cols
                self._create_table(
                    data=contingency_table,
                    header=True,
                    column_widths=(1.2,) + (col_width,) * n_cols,
                    font_size=8.5,
                )

            if normality_tests is not None:
                norm_test_heading = self.document.add_heading(
                    "Tests for Normality", level=4
                )
                norm_test_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # type | p-value | conclusion
                self._create_table(
                    data=normality_tests,
                    header=True,
                    column_widths=(2.2, 1, 2),
                    font_size=8.5,
                    style="Normal Table",
                )

        self.document.add_page_break()

    def _get_bivariate_analysis(self) -> None:
        """Get comparisons and regression-plots for pairs of numeric
        variables.
        """
        bivariate_heading = self.document.add_heading(
            "2. Bivariate Analysis", level=1
        )
        self._format_paragraph_spacing(bivariate_heading, before=0)
        overview_heading = self.document.add_heading("2.1 Overview", level=2)
        self._format_paragraph_spacing(overview_heading)
        self.document.add_picture(
            self.bivariate_graphs["correlation_plot"],
            width=Inches(6.4),
        )
        picture_paragraph = self.document.paragraphs[-1]
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_page_break()

        pairwise_heading = self.document.add_heading(
            "2.2 Regression Plots (Top 20)", level=2
        )
        self._format_paragraph_spacing(pairwise_heading, before=0)
        for idx, var_pair in enumerate(self.bivariate_summaries, start=1):
            heading = self.document.add_heading(
                f"2.2.{idx} {var_pair[0]} vs {var_pair[1]}".title(), level=3
            )
            self._format_paragraph_spacing(heading, before=16, after=5)
            self.document.add_paragraph(self.bivariate_summaries[var_pair])
            self.document.add_picture(
                self.bivariate_graphs["regression_plots"][var_pair],
                width=Inches(3.3),
            )
            picture_paragraph = self.document.paragraphs[-1]
            picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _format_paragraph_spacing(
        self, paragraph: Paragraph, before: int = 15, after: int = 7
    ) -> None:
        """Set the spacing above or below a paragraph.

        Args:
            paragraph (docx.text.paragraph.Paragraph): A paragraph.
            before (int, optional): Size of spacing above the paragraph in pt.
                Defaults to 15.
            after (int, optional): Size of spacing below the paragraph in pt.
                Defaults to 7.
        """
        paragraph.paragraph_format.space_before = Pt(before)
        paragraph.paragraph_format.space_after = Pt(after)

    def _create_table(
        self,
        data: DataFrame,
        column_widths: Sequence = (),
        font_face: str = "Courier New",
        font_size: float = 10,
        style: str = None,
        header: bool = False,
    ) -> None:
        """Generates a table for the supplied ``data``.

        Args:
            data (DataFrame): The data to tabulate.
            column_widths (Sequence, optional): Column dimensions in inches.
                Defaults to ().
            font_face (str, optional): Font for cell text. Defaults to
                "Courier New".
            font_size (float, optional): Font size. Defaults to 10.
            style (str, optional): A `Word` table style. Defaults to
                None.
            header (bool, optional): Whether or not to include column names.
                Defaults to False.
        """
        table = self.document.add_table(
            rows=0,
            cols=len(column_widths),
            style=style or self.document.styles[self.TABLE_STYLE],
        )
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for idx, width in enumerate(column_widths):
            table.columns[idx].width = Inches(width)

        if header:
            cells = table.add_row().cells
            header_labels = [""] + list(data.columns)
            for cell, value in zip(cells, header_labels):
                cell.text = f"{value}"
                # Font size and type-face have to be set at `run` level
                run = cell.paragraphs[0].runs[0]
                run.bold = True
                run.font.size = Pt(font_size)
                run.font.name = font_face

        # Sequentially add and populate rows
        for row_data in data.itertuples():
            cells = table.add_row().cells
            for idx, (cell, value) in enumerate(zip(cells, row_data)):
                try:
                    # Strip trailing zeros from float values
                    text = f"{value:.4f}".rstrip("0").rstrip(".")
                except ValueError:
                    text = f"{value}"

                cell.text = text
                # Font size and type-face have to be set at `run` level
                run = cell.paragraphs[0].runs[0]
                run.font.size = Pt(font_size)
                run.font.name = font_face
                # Make first column values bold if header is True
                if idx == 0 and header:
                    run.bold = True

        # Add empty paragraph. "Spacing" for docx Table isn't yet implemented
        self.document.add_paragraph()

    def _to_file(self) -> None:
        """Save the report as a file."""
        for section in self.document.sections:
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)

        self.document.save(self.OUTPUT_FILENAME)
